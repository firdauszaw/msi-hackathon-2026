"""DOCX-based extractors for TETRA AT command reference material.

This module can pull a best-effort command description mapping and a structured
command-family catalog from the programmer guide.
"""

from typing import Any, Dict, List, Optional
import os
import re
import xml.etree.ElementTree as ET


COMMAND_TOKEN_RE = re.compile(r"^\s*(?P<token>\+?[A-Z0-9]+|&[A-Z])\s*[\u2013-]\s*(?P<title>.+?)\s*$")
DESCRIPTION_OF_COMMANDS_HEADING = "description of at commands"


def _load_docx(path: str):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required to parse docx files: " + str(exc))

    if not os.path.exists(path):
        raise FileNotFoundError(path)

    return Document(path)


def _iter_docx_paragraphs(doc) -> List[Dict[str, Any]]:
    paragraphs: List[Dict[str, Any]] = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        paragraphs.append(
            {
                "index": index,
                "style": style_name,
                "text": text,
            }
        )
    return paragraphs


def _parse_command_heading(text: str) -> Dict[str, Optional[str]]:
    match = COMMAND_TOKEN_RE.match(text)
    if match:
        return {
            "token": match.group("token").upper(),
            "title": match.group("title").strip(),
        }
    return {"token": None, "title": text.strip()}


def extract_mapping_from_docx(path: str) -> Dict[str, str]:
    """Extract a simple mapping of command -> short description from a .docx file.

    Returns an empty dict on failure.
    """
    doc = _load_docx(path)
    texts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            texts.append(t)

    # include table text as well
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    texts.append(t)

    mapping: Dict[str, str] = {}
    combined = "\n".join(texts)

    # Find +COMMAND tokens (letters/numbers, 3-12 chars)
    tokens = set(re.findall(r"\+([A-Z0-9]{3,12})\b", combined, flags=re.I))
    for token in tokens:
        # Find a paragraph containing the token
        pat = re.compile(rf"\+{re.escape(token)}\b", flags=re.I)
        desc = None
        for i, p in enumerate(texts):
            if pat.search(p):
                # remove the token from the paragraph and crop
                desc_candidate = pat.sub("", p).strip(" :-.\n")
                if desc_candidate:
                    desc = desc_candidate
                else:
                    # try the next paragraph as a short description
                    if i + 1 < len(texts):
                        desc = texts[i + 1].strip()
                break
        if desc:
            # Normalize whitespace and limit length
            desc = re.sub(r"\s+", " ", desc)
            if len(desc) > 200:
                desc = desc[:197] + "..."
            mapping[token.upper()] = desc

    return mapping


def extract_command_families_from_docx(path: str) -> List[Dict[str, Any]]:
    """Extract command families and their command headings from a programmer guide."""
    doc = _load_docx(path)
    paragraphs = _iter_docx_paragraphs(doc)

    section_start: Optional[int] = None
    for index, paragraph in enumerate(paragraphs):
        if (
            paragraph["style"].lower().startswith("heading 1")
            and paragraph["text"].strip().lower() == DESCRIPTION_OF_COMMANDS_HEADING
        ):
            section_start = index
            break

    if section_start is None:
        raise ValueError("Could not locate 'Description of AT Commands' heading in the DOCX reference.")

    families: List[Dict[str, Any]] = []
    current_family: Optional[Dict[str, Any]] = None

    for paragraph in paragraphs[section_start + 1 :]:
        style_name = paragraph["style"].lower()
        text = paragraph["text"]

        if style_name.startswith("heading 1"):
            break

        if style_name.startswith("heading 2"):
            current_family = {
                "name": text,
                "paragraph_index": paragraph["index"],
                "commands": [],
            }
            families.append(current_family)
            continue

        if style_name.startswith("heading 3") and current_family is not None:
            parsed_heading = _parse_command_heading(text)
            current_family["commands"].append(
                {
                    "token": parsed_heading["token"],
                    "title": parsed_heading["title"],
                    "heading": text,
                    "paragraph_index": paragraph["index"],
                }
            )

    return families


def write_command_families_xml(docx_path: str, output_path: str) -> List[Dict[str, Any]]:
    """Extract command families from DOCX and write them as XML."""
    families = extract_command_families_from_docx(docx_path)

    root = ET.Element(
        "tetra-at-command-families",
        {
            "version": "1.0",
            "source-doc": os.path.basename(docx_path),
            "family-count": str(len(families)),
        },
    )

    total_commands = 0
    for family in families:
        commands = family["commands"]
        total_commands += len(commands)
        family_element = ET.SubElement(
            root,
            "family",
            {
                "name": family["name"],
                "paragraph-index": str(family["paragraph_index"]),
                "command-count": str(len(commands)),
            },
        )
        for command in commands:
            attrs = {
                "paragraph-index": str(command["paragraph_index"]),
                "heading": command["heading"],
            }
            if command["token"]:
                attrs["token"] = command["token"]
            command_element = ET.SubElement(family_element, "command", attrs)
            command_element.text = command["title"]

    root.set("command-count", str(total_commands))
    tree = ET.ElementTree(root)
    ET.indent(tree, space="  ")
    tree.write(output_path, encoding="utf-8", xml_declaration=True)
    return families


def safe_load_docx_mapping(path: str) -> Dict[str, str]:
    try:
        return extract_mapping_from_docx(path)
    except Exception:
        return {}


if __name__ == "__main__":
    print("docparser: use safe_load_docx_mapping(path) or write_command_families_xml(docx_path, output_path).")
